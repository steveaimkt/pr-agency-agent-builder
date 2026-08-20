#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
산출물을 워드 파일로 만든다.

  python3 tools/md2docx.py output/03-월간보고서.md
  python3 tools/md2docx.py output/                 폴더 안 md 전부

같은 자리에 .docx 가 생긴다. 고객사가 열어서 고칠 수 있다.
색과 글꼴은 tools/design.md 를 따른다. 새로 고르지 않는다.

워드 뼈대로 짠다 · outputs.md 「워드로 갈 때」 와 같다
  표지    맨 처음 # 제목과 바로 뒤 > 줄을 표지로 올린다
  목차    장이 셋 넘으면 자동으로 넣는다 (워드에서 F9 로 채워진다)
  본문    Heading 1·2·3 스타일을 쓴다. 그래야 목차에 잡힌다
  꼬리말  쪽번호

읽는 것
  # ## ###        제목
  | 표 |          워드 표로. 가로선만. 머리행은 작고 흐리게
  ```숫자 ...```   숫자 카드
  ```막대 ...```   막대 (칸 문자로 그린다. 워드에서 안 깨진다)
  ```달성 ...```   달성률 (100 넘으면 초록, 90 미만 빨강)
  [FACT-CHECK]    형광 표시
  [확인 필요]      형광 표시
"""
import sys, os, re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print('python-docx 가 없습니다 · pip3 install python-docx'); sys.exit(1)

# ─── 색 · tools/design.md 와 같은 값 ───
INK   = RGBColor(0x2A, 0x2A, 0x28)
BODY  = RGBColor(0x4C, 0x4D, 0x59)
META  = RGBColor(0x84, 0x83, 0x83)
SAND  = RGBColor(0xC9, 0xB0, 0xA0)
GOOD  = RGBColor(0x37, 0x62, 0x4A)
BAD   = RGBColor(0xA8, 0x32, 0x26)
FONT  = 'Pretendard'
LINE_HEX, SOFT_HEX, HL_HEX = 'E6E2DC', 'F0EDE8', 'EBFF2C'

MARKS = ('[FACT-CHECK]', '[확인 필요]')
ROW_RE = re.compile(r'^(.+?)\s+([\d,]+(?:\.\d+)?)\s*([%건명원개점회편장]*)\s*$')
KPI_RE = re.compile(r'^([\d,]+(?:\.\d+)?[%가-힣]*)\s+(.+?)(?:\s*·\s*(좋음|나쁨))?\s*$')


def _shade(cell, hex_color):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(el)


def _borders_horizontal_only(table):
    """세로줄을 긋지 않는다. 가로선만."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom', 'insideH'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), LINE_HEX)
        borders.append(e)
    for edge in ('left', 'right', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none'); e.set(qn('w:sz'), '0')
        borders.append(e)
    tblPr.append(borders)


def _run(par, text, size=10.5, bold=False, color=BODY, mark=False):
    r = par.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if mark:
        hl = OxmlElement('w:shd')
        hl.set(qn('w:val'), 'clear'); hl.set(qn('w:fill'), HL_HEX)
        r._element.get_or_add_rPr().append(hl)
    return r


def _text_par(doc, raw, size=10.5, color=BODY, space_after=6):
    """볼드와 확인 표시를 살려서 한 문단을 쓴다."""
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.line_spacing = 1.5
    # **볼드** 와 확인 표시를 조각으로 가른다
    parts = re.split(r'(\*\*[^*]+\*\*|\[FACT-CHECK\]|\[확인 필요\])', raw)
    for s in parts:
        if not s:
            continue
        if s in MARKS:
            _run(par, s, size, True, INK, mark=True)
        elif s.startswith('**') and s.endswith('**'):
            _run(par, s[2:-2], size, True, INK)
        else:
            _run(par, s, size, False, color)
    return par


def _kpi_cards(doc, title, lines):
    rows = []
    for ln in lines:
        m = KPI_RE.match(ln.strip())
        if m:
            rows.append((m.group(1), m.group(2), m.group(3)))
    if not rows:
        return
    if title:
        _text_par(doc, title, 9.5, META, 3)
    t = doc.add_table(rows=2, cols=len(rows))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _borders_horizontal_only(t)
    for i, (big, label, judge) in enumerate(rows):
        c = t.cell(0, i); c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        col = GOOD if judge == '좋음' else BAD if judge == '나쁨' else INK
        _run(p, big, 20, True, col)
        _shade(c, SOFT_HEX)
        c2 = t.cell(1, i); c2.text = ''
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, label, 9, False, META)
        _shade(c2, SOFT_HEX)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _bars(doc, title, lines, rate=False):
    rows = []
    for ln in lines:
        m = ROW_RE.match(ln.strip())
        if m:
            rows.append((m.group(1).strip(), float(m.group(2).replace(',', '')), m.group(3)))
    if len(rows) < 2:
        return
    if title:
        _text_par(doc, title, 10, INK, 3).runs[0].font.bold = True
    mx = max(r[1] for r in rows) or 1
    t = doc.add_table(rows=len(rows), cols=3)
    _borders_horizontal_only(t)
    for i, (label, val, unit) in enumerate(rows):
        c0 = t.cell(i, 0); c0.text = ''
        _run(c0.paragraphs[0], label, 10, False, BODY)
        # 막대는 칸 문자로 그린다. 워드 어디서 열어도 안 깨진다
        n = max(1, round(28 * (val / mx)))
        c1 = t.cell(i, 1); c1.text = ''
        if rate:
            col = GOOD if val >= 100 else BAD if val < 90 else SAND
        else:
            col = SAND if val < mx else RGBColor(0xA8, 0x8B, 0x78)
        _run(c1.paragraphs[0], '█' * n, 9, False, col)
        c2 = t.cell(i, 2); c2.text = ''
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        txt = f'{val:g}{unit}' + ('%' if rate and not unit else '')
        _run(p2, txt, 10, val >= mx, INK)
    for row in t.rows:
        row.cells[0].width = Cm(4.2); row.cells[1].width = Cm(7.5); row.cells[2].width = Cm(2.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _table(doc, rows):
    """마크다운 표 → 워드 표. 세로줄 없이 가로선만."""
    head, body = rows[0], rows[1:]
    t = doc.add_table(rows=len(rows), cols=len(head))
    _borders_horizontal_only(t)
    for j, cell in enumerate(head):
        c = t.cell(0, j); c.text = ''
        _run(c.paragraphs[0], re.sub(r'\*\*', '', cell), 9, True, META)
        _shade(c, SOFT_HEX)
    for i, r in enumerate(body, start=1):
        for j in range(len(head)):
            cell = r[j] if j < len(r) else ''
            c = t.cell(i, j); c.text = ''
            p = c.paragraphs[0]
            is_num = bool(re.fullmatch(r'[\d,.\-+%건명원개점회편장억만 ]+', cell.strip() or 'x'))
            if is_num and cell.strip():
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for s in re.split(r'(\*\*[^*]+\*\*|\[FACT-CHECK\]|\[확인 필요\])', cell):
                if not s:
                    continue
                if s in MARKS:
                    _run(p, s, 9.5, True, INK, mark=True)
                elif s.startswith('**') and s.endswith('**'):
                    _run(p, s[2:-2], 9.5, True, INK)
                else:
                    _run(p, s, 9.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _field(par, code, placeholder=''):
    """워드 필드를 넣는다. 쪽번호에만 쓴다."""
    r = par.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = code
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    tx = OxmlElement('w:t'); tx.text = placeholder
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for el in (b, i, s, tx, e):
        r._r.append(el)
    r.font.name = FONT; r.font.size = Pt(9); r.font.color.rgb = META
    return r


def _cover(doc, title, subs):
    """표지 · 제목 하나와 밑에 몇 줄"""
    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(18)
    _run(par, title, 26, True, INK)
    for s in subs[:4]:
        p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(4)
        _run(p2, s, 11, False, META)
    doc.add_page_break()


def _toc(doc, heads):
    """차례 · 쪽번호를 붙이지 않는다.

    워드 자동 목차 필드는 프로그램마다 탭이 다르게 잡혀
    쪽번호가 줄바꿈으로 떨어진다. 그래서 장 이름만 적는다.
    제목이 Heading 스타일이라 워드 탐색 창으로 넘어다닐 수 있다.
    """
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(12)
    _run(par, '차례', 15, True, INK)
    for lvl, txt in heads:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.left_indent = Cm(0 if lvl <= 2 else 0.7)
        _run(p2, txt, 11 if lvl <= 2 else 10, lvl <= 2, INK if lvl <= 2 else BODY)
    p3 = doc.add_paragraph(); p3.paragraph_format.space_before = Pt(14)
    _run(p3, '쪽번호는 아래쪽에 있습니다. 워드 탐색 창에서 장을 눌러 넘어갈 수 있습니다.', 9, False, META)
    doc.add_page_break()


def _page_number(doc):
    ft = doc.sections[0].footer.paragraphs[0]
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(ft, 'PAGE', '')


def _heading(doc, level, text):
    """Heading 스타일을 쓰되 색과 글꼴은 우리 것으로 덮는다"""
    style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}[level]
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
    par.paragraph_format.space_after = Pt(6)
    size = {1: 18, 2: 14, 3: 12, 4: 11}[level]
    _run(par, text, size, True, INK)
    return par


def convert(src, dst):
    lines = open(src, encoding='utf-8').read().split('\n')
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.5)
    st = doc.styles['Normal']
    st.font.name = FONT; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # ── 표지와 목차 ──
    body_start, cover_title, cover_subs = 0, '', []
    for k, ln in enumerate(lines):
        if ln.startswith('# '):
            cover_title = re.sub(r'\*\*', '', ln[2:]).strip()
            body_start = k + 1
            break
    if cover_title:
        for ln in lines[body_start:body_start + 8]:
            s = ln.strip()
            if s.startswith('>'):
                cover_subs.append(re.sub(r'\*\*|`', '', s.lstrip('>').strip()))
            elif s and not s.startswith('#'):
                break
        _cover(doc, cover_title, cover_subs)
        lines = lines[body_start:]
    _page_number(doc)

    heads = []
    for x in lines:
        mm = re.match(r'^(#{2,3})\s+(.+)$', x)
        if mm:
            heads.append((len(mm.group(1)), re.sub(r'\*\*|`', '', mm.group(2)).strip()))
    if len(heads) >= 3:
        _toc(doc, heads)

    i, n = 0, len(lines)
    while i < n:
        ln = lines[i]

        # 도식 블록
        m = re.match(r'^```(숫자|막대|달성)\s*(.*)$', ln.strip())
        if m:
            kind, title = m.group(1), m.group(2).strip()
            body = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                body.append(lines[i]); i += 1
            i += 1
            if kind == '숫자':
                _kpi_cards(doc, title, body)
            else:
                _bars(doc, title, body, rate=(kind == '달성'))
            continue

        # 그 밖의 코드 블록은 그대로 옮긴다
        if ln.strip().startswith('```'):
            body = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                body.append(lines[i]); i += 1
            i += 1
            if any(s.strip() for s in body):
                par = doc.add_paragraph()
                par.paragraph_format.space_after = Pt(8)
                r = par.add_run('\n'.join(body))
                r.font.name = 'Menlo'; r.font.size = Pt(9); r.font.color.rgb = BODY
                _shade_par = OxmlElement('w:shd')
                _shade_par.set(qn('w:val'), 'clear'); _shade_par.set(qn('w:fill'), SOFT_HEX)
                par._p.get_or_add_pPr().append(_shade_par)
            continue

        # 표
        if ln.strip().startswith('|') and i + 1 < n and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not re.match(r'^[\s:|-]+$', ''.join(cells)):
                    rows.append(cells)
                i += 1
            if rows:
                _table(doc, rows)
            continue

        # 제목
        m = re.match(r'^(#{1,4})\s+(.+)$', ln)
        if m:
            lvl, txt = len(m.group(1)), re.sub(r'\*\*', '', m.group(2)).strip()
            _heading(doc, lvl, txt)
            i += 1
            continue

        # 인용
        if ln.strip().startswith('>'):
            _text_par(doc, ln.strip().lstrip('>').strip(), 10, META, 6)
            i += 1
            continue

        # 목록
        m = re.match(r'^\s*[-*]\s+(.+)$', ln)
        if m:
            par = _text_par(doc, m.group(1), 10.5, BODY, 3)
            par.paragraph_format.left_indent = Cm(0.6)
            i += 1
            continue

        # 가로선은 건너뛴다
        if re.match(r'^\s*---+\s*$', ln):
            i += 1
            continue

        if ln.strip():
            _text_par(doc, ln.strip())
        i += 1

    doc.save(dst)


def main():
    if len(sys.argv) < 2:
        print(__doc__); sys.exit(1)
    target = sys.argv[1]
    files = []
    if os.path.isdir(target):
        files = [os.path.join(target, f) for f in sorted(os.listdir(target)) if f.endswith('.md')]
    elif target.endswith('.md'):
        files = [target]
    if not files:
        print(f'md 파일을 못 찾았습니다: {target}'); sys.exit(1)
    for f in files:
        out = f[:-3] + '.docx'
        convert(f, out)
        print(f'  {os.path.basename(out)}')
    print(f'\n{len(files)}개 만들었습니다. 워드에서 열어 그대로 고칠 수 있습니다.')


if __name__ == '__main__':
    main()
